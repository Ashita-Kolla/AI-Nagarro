import os
import json
import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from core.llm_utils import call_llm, parse_json_from_llm

def generate_brd(ba_output: dict, project_name: str):
    try:
        out_dir = os.path.join("outputs", "BA")
        os.makedirs(out_dir, exist_ok=True)
        doc = Document()
        
        style_h1 = doc.styles['Heading 1']
        font_h1 = style_h1.font
        font_h1.name = 'Calibri'
        font_h1.size = Pt(16)
        font_h1.bold = True
        font_h1.color.rgb = RGBColor(0x1F, 0x38, 0x64)
        
        style_h2 = doc.styles['Heading 2']
        font_h2 = style_h2.font
        font_h2.name = 'Calibri'
        font_h2.size = Pt(13)
        font_h2.bold = True
        font_h2.color.rgb = RGBColor(0x2F, 0x2F, 0x2F)
        
        style_normal = doc.styles['Normal']
        font_normal = style_normal.font
        font_normal.name = 'Calibri'
        font_normal.size = Pt(11)
        font_normal.color.rgb = RGBColor(0x00, 0x00, 0x00)
        
        doc.add_paragraph("[Company Logo Placeholder]")
        
        title = doc.add_paragraph("Business Requirements Document")
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.runs[0].font.size = Pt(24)
        title.runs[0].bold = True
        
        doc.add_paragraph(f"Project: {project_name}")
        doc.add_paragraph(f"Date: {datetime.date.today().strftime('%Y-%m-%d')}")
        doc.add_paragraph("Version: 1.0")
        doc.add_paragraph("Status: Draft")
        
        doc.add_page_break()
        
        doc.add_heading("Table of Contents", level=1)
        doc.add_paragraph("1. Executive Summary")
        doc.add_paragraph("2. Business Requirements")
        doc.add_paragraph("3. User Stories")
        doc.add_paragraph("4. Functional Requirements")
        doc.add_paragraph("5. Non-Functional Requirements")
        doc.add_paragraph("6. Assumptions")
        doc.add_paragraph("7. Out of Scope")
        doc.add_paragraph("8. Document History")
        
        doc.add_page_break()
        
        doc.add_heading("1. Executive Summary", level=1)
        exec_summary_prompt = f"""
Based on the following approved Business Analysis output, write a 3-4 sentence executive summary 
summarising what is being built, who it is for, and what the key goals are.
Write in plain business English.

BA OUTPUT:
{json.dumps(ba_output)}
"""
        exec_summary = call_llm(exec_summary_prompt) or "Executive summary could not be generated."
        doc.add_paragraph(exec_summary)
        
        doc.add_heading("2. Business Requirements", level=1)
        for i, br in enumerate(ba_output.get("business_requirements", []), 1):
            doc.add_paragraph(br, style='List Number')
            
        doc.add_heading("3. User Stories", level=1)
        for us in ba_output.get("user_stories", []):
            doc.add_heading(f"{us.get('id', '')} — {us.get('action', '')}", level=2)
            doc.add_paragraph(f"As a {us.get('role', '')}, I want {us.get('action', '')}, so that {us.get('benefit', '')}")
            
            table = doc.add_table(rows=1, cols=1)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_run = hdr_cells[0].paragraphs[0].add_run('Acceptance Criteria')
            hdr_run.bold = True
            
            for ac in us.get("acceptance_criteria", []):
                row_cells = table.add_row().cells
                row_cells[0].text = ac
                
        doc.add_heading("4. Functional Requirements", level=1)
        for fr in ba_output.get("functional_requirements", []):
            doc.add_paragraph(fr, style='List Number')
            
        doc.add_heading("5. Non-Functional Requirements", level=1)
        nfrs = ba_output.get("non_functional_requirements", {})
        for category, reqs in nfrs.items():
            if reqs:
                doc.add_heading(category.capitalize(), level=2)
                for req in reqs:
                    doc.add_paragraph(req, style='List Bullet')
                    
        doc.add_heading("6. Assumptions", level=1)
        for assump in ba_output.get("assumptions", []):
            doc.add_paragraph(assump, style='List Number')
            
        doc.add_heading("7. Out of Scope", level=1)
        for oos in ba_output.get("out_of_scope", []):
            doc.add_paragraph(oos, style='List Number')
            
        doc.add_heading("8. Document History", level=1)
        hist_table = doc.add_table(rows=2, cols=4)
        hist_table.style = 'Table Grid'
        
        headers = ['Version', 'Date', 'Author', 'Notes']
        for i, header in enumerate(headers):
            cell = hist_table.cell(0, i)
            cell.text = header
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
            
        row_cells = hist_table.rows[1].cells
        row_cells[0].text = "1.0"
        row_cells[1].text = datetime.date.today().strftime("%Y-%m-%d")
        row_cells[2].text = "ARIA BA Agent"
        row_cells[3].text = "Initial draft"
        
        brd_path = os.path.join(out_dir, "BRD.docx")
        doc.save(brd_path)
        print(f"Successfully generated {brd_path}")
        return brd_path
    except Exception as e:
        print(f"Error generating BRD.docx: {e}")
        return None

def post_approval(data: dict, context_manager):
    """Hook called by the agent runner after human approval."""
    context = context_manager.get_context()
    supervisor_output = context.get("Supervisor", {})
    project_name = "Unknown Project"
    if isinstance(supervisor_output, dict):
        project_name = supervisor_output.get("project_name", "Unknown Project")
    return generate_brd(data, project_name)

def run(context_manager, correction: str = None) -> dict:
    """
    Runs the BA agent.
    Returns the parsed JSON output.
    """
    context = context_manager.get_context()
    user_brief = context.get("USER_BRIEF", "")
    supervisor_output = context.get("Supervisor", {})
    
    prompt_template = """
You are the BA Agent for ARIA, a multi-agent SDLC 
system. You perform deep business analysis on a 
project brief and produce structured, professional 
business analysis artifacts.

SUPERVISOR OUTPUT:
{supervisor_output}

USER BRIEF:
{user_brief}
{human_correction}
---

YOUR JOB:

Produce a complete business analysis. Be specific. 
Be thorough. Do not be vague. Every item you produce 
will be used by 6 downstream agents — Architect, 
Developer, QA, DevOps, PM, and Optimisation.

1. BUSINESS REQUIREMENTS
   Write 5-7 high-level business requirements.
   Format: "BR-001: The system must..."
   These are business needs, not technical specs.
   Cover: core purpose, user needs, performance, security, scalability.

2. USER STORIES
   Write one user story per major feature.
   Minimum 4 stories, maximum 6.
   
   Format exactly:
   - id: US-001
   - role: specific type of user
   - action: what they want to do (be concise)
   - benefit: measurable outcome
   - acceptance_criteria: list of 2-3 measurable, testable conditions.

3. FUNCTIONAL REQUIREMENTS
   One requirement per major system behaviour. Maximum 8.
   Format: "FR-001: The system must..."

4. NON-FUNCTIONAL REQUIREMENTS
   Cover only the most relevant categories (max 2 items each):
   - Performance, Security, Scalability, Availability

5. ASSUMPTIONS
   List up to 5 key assumptions.

6. OUT OF SCOPE
   List up to 4 explicitly excluded items.

7. CONFIDENCE SCORE
   Score 0-100 with one sentence of reasoning.

---

CRITICAL RULES:
- Output ONLY raw valid JSON. No markdown. 
  No backticks. No prose before or after.
- Do not invent features not in the brief or 
  implied by it.
- Every acceptance criterion must be testable 
  by a QA engineer without asking questions.
- If the brief mentions compliance (GDPR, HIPAA, 
  PCI-DSS, WCAG), create dedicated requirements 
  for it — do not just mention it in passing.

---

OUTPUT FORMAT:

{{
  "business_requirements": [
    "BR-001: ...",
    "BR-002: ..."
  ],
  "user_stories": [
    {{
      "id": "US-001",
      "role": "",
      "action": "",
      "benefit": "",
      "acceptance_criteria": [
        "criterion 1",
        "criterion 2",
        "criterion 3"
      ]
    }}
  ],
  "functional_requirements": [
    "FR-001: ...",
    "FR-002: ..."
  ],
  "non_functional_requirements": {{
    "performance": [],
    "security": [],
    "scalability": [],
    "accessibility": [],
    "availability": [],
    "compatibility": []
  }},
  "assumptions": [
    "Assumption 1: ..."
  ],
  "out_of_scope": [
    "Feature or capability not included"
  ],
  "confidence_score": 0,
  "confidence_reasoning": ""
}}
"""

    human_correction_text = f"\nHUMAN CORRECTION: {correction}\n" if correction else ""
    
    prompt = prompt_template.format(
        supervisor_output=json.dumps(supervisor_output) if isinstance(supervisor_output, dict) else supervisor_output,
        user_brief=user_brief,
        human_correction=human_correction_text
    )
    
    print("Calling LLM for BA analysis...")
    response_text = call_llm(prompt, agent_name='BA')
    data = parse_json_from_llm(response_text)
    
    if not data:
        print("Error: LLM returned invalid JSON.")
        return None
        
    return data
