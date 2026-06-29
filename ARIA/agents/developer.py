import os
import json
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from core.llm_utils import call_llm, parse_json_from_llm

from core.context_compressor import compress_context_for_agent

def build_foundation_prompt(tech_stack: dict, full_ctx_str: str, correction: str = None) -> str:
    frontend = tech_stack.get("frontend", "Not specified")
    backend = tech_stack.get("backend", "Not specified")
    database = tech_stack.get("database", "Not specified")
    
    prompt = f"""You are the Developer Agent for ARIA. Output ONLY a single valid JSON object. Nothing else.

⚠️ ABSOLUTE OUTPUT RULES:
1. Do NOT write any text before or after the JSON.
2. Do NOT wrap the JSON in markdown fences (no ```json).
3. ALL code inside JSON string values MUST use \\n for newlines.
4. ALL double quotes inside code strings MUST be escaped as \\".

PROJECT CONTEXT:
{full_ctx_str}

Generate ONLY the project foundation files for a {frontend} + {backend} project.

Foundation files only:
- package.json or requirements.txt
- main entry point (main.py or App.jsx)
- config files (.env.example, tailwind.config.js etc)
- database connection or storage utility

Do NOT over-engineer the foundation, BUT ensure standard architecture separation. Put frontend code inside a `frontend/` directory (e.g. `frontend/package.json`) and backend code inside a `backend/` directory (e.g. `backend/main.py`). Skip unnecessary boilerplate.

CODE REQUIREMENTS:
1. FRONTEND: Use {frontend}. Follow standard conventions. Avoid complex build steps if possible.
2. BACKEND: Use {backend}. Ensure the application can serve both the API and the static frontend.
3. DATABASE: Use {database}. Set up the initial connection logic.

"""
    if correction:
        prompt += f"\n=== HUMAN/QA CORRECTION ===\n{correction}\n"
        prompt += "\n⚠️ Keep your fix precise. Output the full updated codebase for the foundation.\n"

    prompt += """
OUTPUT FORMAT:
{{
  "files": {{"filepath": "content"}},
  "setup_instructions": {{"steps": [], "environment_variables": []}}
}}
"""
    return prompt

def build_epic_prompt(epic: dict, tech_stack: dict, existing_files: list, user_brief: str = "") -> str:
    frontend = tech_stack.get("frontend", "Not specified")
    backend = tech_stack.get("backend", "Not specified")
    database = tech_stack.get("database", "Not specified")
    
    prompt = f"""You are the Developer Agent for ARIA. Output ONLY a single valid JSON object. Nothing else.

⚠️ ABSOLUTE OUTPUT RULES:
1. Do NOT write any text before or after the JSON.
2. Do NOT wrap the JSON in markdown fences (no ```json).
3. ALL code inside JSON string values MUST use \\n for newlines.
4. ALL double quotes inside code strings MUST be escaped as \\".

PROJECT GOAL:
{user_brief}

Generate code files for this epic ONLY:
Epic: {epic.get('title', 'Unknown')}
Tasks: {json.dumps(epic.get('tasks', []))}
Tech stack: {frontend} + {backend} + {database}

Already generated files (DO NOT regenerate these unless modifying them for this epic):
{chr(10).join(existing_files)}

Generate ONLY the new/modified files needed for this epic.
Consolidate your code into a minimal number of files, BUT ensure standard architecture separation.
Put frontend code inside a `frontend/` directory (e.g. `frontend/package.json`, `frontend/src/App.jsx`) and backend code inside a `backend/` directory (e.g. `backend/main.py`).
If modifying an existing file, provide the complete, updated file. Ensure all code is strictly functional and fully complete. No placeholders. No TODOs.

OUTPUT FORMAT:
{{
  "files": {{"filepath": "complete file content"}}
}}
"""
    return prompt

def build_patch_prompt(tech_stack: dict, existing_files: list, correction: str, user_brief: str = "") -> str:
    prompt = f"""You are the Developer Agent for ARIA. Output ONLY a single valid JSON object. Nothing else.

⚠️ ABSOLUTE OUTPUT RULES:
1. Do NOT write any text before or after the JSON.
2. Do NOT wrap the JSON in markdown fences (no ```json).
3. ALL code inside JSON string values MUST use \\n for newlines.
4. ALL double quotes inside code strings MUST be escaped as \\".

PROJECT GOAL:
{user_brief}

You are in PATCH MODE. An automated QA run or a human reviewer has found issues with the codebase.
Feedback/Test Logs:
{correction}

Already generated files (DO NOT regenerate these unless modifying them to fix the issues):
{chr(10).join(existing_files)}

Generate ONLY the new/modified files needed to fix the issues mentioned in the feedback.
Consolidate your code into a minimal number of files, BUT ensure standard architecture separation.
Put frontend code inside a `frontend/` directory (e.g. `frontend/package.json`, `frontend/src/App.jsx`) and backend code inside a `backend/` directory (e.g. `backend/main.py`).
If modifying an existing file, provide the complete, updated file. Ensure all code is strictly functional and fully complete. No placeholders. No TODOs.

OUTPUT FORMAT:
{{
  "files": {{"filepath": "complete file content"}}
}}
"""
    return prompt

def run(context_manager, correction: str = None) -> dict:
    context = context_manager.get_context()
    
    if not context.get("Architect") or not context.get("Planner"):
        print("[Developer Agent Error] Missing Architect or Planner output.")
        return {}

    # Compress Context
    compressed = compress_context_for_agent("Developer", context)
    full_ctx_str = json.dumps(compressed, indent=2)

    tech_stack = context["Architect"].get("technology_stack", {})
    epics = context["Planner"].get("epics", [])
    user_brief = context.get("USER_BRIEF", "")
    
    all_files = {}
    setup = {}

    try:
        # Patch Mode: if there's a correction and Developer has already run
        if correction and "Developer" in context:
            print("[Developer] Patch Mode Activated! Processing corrections...")
            existing_files_dict = context["Developer"].get("files", {})
            existing_files = list(existing_files_dict.keys())
            
            patch_prompt = build_patch_prompt(tech_stack, existing_files, correction, user_brief)
            patch_res = call_llm(patch_prompt, agent_name="Developer", max_tokens=4000)
            parsed_patch = parse_json_from_llm(patch_res)
            
            if parsed_patch:
                patched_files = parsed_patch.get("files", {})
                existing_files_dict.update(patched_files) # Merge patches
                return {"files": existing_files_dict, "setup_instructions": context["Developer"].get("setup_instructions", {})}
            else:
                print("[Developer] Failed to parse patch output.")
                return context["Developer"]
                
        # Chunk 1: Foundation
        print("[Developer] Generating foundation...")
        foundation_prompt = build_foundation_prompt(tech_stack, full_ctx_str, correction)
        foundation_res = call_llm(foundation_prompt, agent_name="Developer", max_tokens=4000)
        parsed_foundation = parse_json_from_llm(foundation_res)
        
        if parsed_foundation:
            all_files.update(parsed_foundation.get("files", {}))
            setup = parsed_foundation.get("setup_instructions", {})
            
        # Chunk 2-N: Epics
        for idx, epic in enumerate(epics):
            print(f"[Developer] Generating Epic {idx+1}/{len(epics)}: {epic.get('title')}")
            epic_prompt = build_epic_prompt(epic, tech_stack, list(all_files.keys()), user_brief=user_brief)
            epic_res = call_llm(epic_prompt, agent_name="Developer", max_tokens=4000)
            parsed_epic = parse_json_from_llm(epic_res)
            
            if parsed_epic:
                # Merge new/updated files
                for fpath, content in parsed_epic.get("files", {}).items():
                    all_files[fpath] = content

        return {"files": all_files, "setup_instructions": setup}
        
    except Exception as e:
        print(f"Developer Agent Error: {str(e)}")
        return {}

def post_approval(data: dict, context_manager) -> list:
    out_dir = os.path.join("outputs", "Developer")
    codebase_dir = os.path.join(out_dir, "codebase")
    os.makedirs(codebase_dir, exist_ok=True)
    
    generated_files = []

    # 1. Save Full Output JSON (exclude files to save space)
    full_json_path = os.path.join(out_dir, "developer_summary.json")
    with open(full_json_path, "w", encoding="utf-8") as f:
        summary_data = {k: v for k, v in data.items() if k != "files"}
        json.dump(summary_data, f, indent=2, ensure_ascii=False)
    generated_files.append(full_json_path)

    # 2. Extract and physically save the codebase files
    files = data.get("files", {})
    failed_writes = []
    
    for file_path, content in files.items():
        if file_path and content:
            safe_path = os.path.normpath(file_path).lstrip("\\/")
            if ".." in safe_path:
                continue
                
            full_path = os.path.join(codebase_dir, safe_path)
            try:
                os.makedirs(os.path.dirname(full_path), exist_ok=True)
                with open(full_path, "w", encoding="utf-8") as f:
                    f.write(content)
                generated_files.append(full_path)
            except Exception as e:
                failed_writes.append(f"{safe_path}: {e}")

    if failed_writes:
        print(f"✗ Failed to write {len(failed_writes)} files:")
        for err in failed_writes:
            print(f"  {err}")

    # 3. Generate .docx document using python-docx
    docx_path = os.path.join(out_dir, "developer_tasks.docx")
    try:
        generate_developer_tasks_docx(docx_path, context_manager, data)
        generated_files.append(docx_path)
    except Exception as e:
        print(f"Failed to generate developer_tasks.docx: {e}")

    print(f"Developer artifacts exported to {out_dir}.")
    return generated_files

def generate_developer_tasks_docx(docx_path, context_manager, dev_data):
    doc = Document()
    planner_data = context_manager.get_summary("Planner") or {}
    supervisor_data = context_manager.get_summary("Supervisor") or {}
    project_name = supervisor_data.get("project_name", "Unknown Project")
    summary = planner_data.get("summary", {})

    # Styles matching BRD
    style_h1 = doc.styles['Heading 1']
    font_h1 = style_h1.font
    font_h1.name = 'Calibri'
    font_h1.size = Pt(16)
    font_h1.bold = True
    font_h1.color.rgb = RGBColor(0x1F, 0x38, 0x64)
    
    style_h2 = doc.styles['Heading 2']
    font_h2 = style_h2.font
    font_h2.name = 'Calibri'
    font_h2.size = Pt(13)
    font_h2.bold = True
    font_h2.color.rgb = RGBColor(0x2F, 0x2F, 0x2F)
    
    style_normal = doc.styles['Normal']
    font_normal = style_normal.font
    font_normal.name = 'Calibri'
    font_normal.size = Pt(11)

    # Cover Page
    title = doc.add_paragraph("Developer Task Breakdown")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.size = Pt(24)
    title.runs[0].bold = True

    doc.add_paragraph(f"Project: {project_name}")
    import datetime
    doc.add_paragraph(f"Date: {datetime.date.today().strftime('%Y-%m-%d')}")
    doc.add_paragraph(f"Total Story Points: {summary.get('total_story_points', 0)}")
    doc.add_paragraph(f"Estimated Duration: {summary.get('estimated_developer_days', 0)} days")
    doc.add_page_break()

    # 1. Executive Summary
    doc.add_heading("1. Executive Summary", level=1)
    doc.add_paragraph(f"This document outlines the development tasks required to complete {project_name}. "
                      f"It includes a total of {summary.get('total_story_points', 0)} story points spanning "
                      f"across {len(planner_data.get('epics', []))} epics. The recommended team size is "
                      f"{summary.get('recommended_team_size', 1)} developers.")

    # 2. Sprint Plan
    doc.add_heading("2. Sprint Plan", level=1)
    for sprint in summary.get("sprint_breakdown", []):
        doc.add_heading(f"Sprint {sprint.get('sprint', 1)}", level=2)
        doc.add_paragraph(f"Focus: {sprint.get('focus', '')} ({sprint.get('story_points', 0)} pts)")
        
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Epic ID'
        hdr_cells[1].text = 'Task ID'
        hdr_cells[2].text = 'Story Points'
        hdr_cells[3].text = 'Complexity'
        
        for epic_id in sprint.get('epics', []):
            for e in planner_data.get('epics', []):
                if e.get('id') == epic_id:
                    for t in e.get('tasks', []):
                        row_cells = table.add_row().cells
                        row_cells[0].text = epic_id
                        row_cells[1].text = t.get('id', '')
                        row_cells[2].text = str(t.get('story_points', ''))
                        row_cells[3].text = t.get('complexity', '')

    # 3. Epic Breakdown
    doc.add_heading("3. Epic Breakdown", level=1)
    for epic in planner_data.get("epics", []):
        doc.add_heading(f"{epic.get('id', '')}: {epic.get('title', '')}", level=2)
        doc.add_paragraph(f"Technical Notes: {epic.get('technical_notes', '')}")
        
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Task ID'
        hdr_cells[1].text = 'Title'
        hdr_cells[2].text = 'Points'
        hdr_cells[3].text = 'Tech Area'
        
        for task in epic.get("tasks", []):
            row_cells = table.add_row().cells
            row_cells[0].text = task.get("id", "")
            row_cells[1].text = task.get("title", "")
            row_cells[2].text = str(task.get("story_points", ""))
            row_cells[3].text = task.get("tech_area", "")

    # 4. Setup Instructions
    doc.add_heading("4. Setup Instructions", level=1)
    setup = dev_data.get("setup_instructions", {})
    for step in setup.get("steps", []):
        doc.add_paragraph(step, style='List Number')

    # 5. Environment Variables
    doc.add_heading("5. Environment Variables", level=1)
    envs = setup.get("environment_variables", [])
    if envs:
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Variable'
        hdr_cells[1].text = 'Description'
        hdr_cells[2].text = 'Example'
        
        for env in envs:
            row_cells = table.add_row().cells
            if isinstance(env, dict):
                # Expected format: {"key": ..., "description": ..., "example": ...}
                row_cells[0].text = env.get("key", "")
                row_cells[1].text = env.get("description", "")
                row_cells[2].text = env.get("example", "")
            else:
                # Flat string format: "DATABASE_URL=sqlite:///tasks.db"
                env_str = str(env)
                parts = env_str.split("=", 1)
                row_cells[0].text = parts[0].strip()
                row_cells[1].text = ""
                row_cells[2].text = parts[1].strip() if len(parts) > 1 else env_str

    doc.save(docx_path)
