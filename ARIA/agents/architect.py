import os
import json
import datetime
from abc import ABC, abstractmethod
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from core.llm_utils import call_llm, parse_json_from_llm

# ==========================================
# EXPORTER FRAMEWORK
# ==========================================

class BaseExporter(ABC):
    """Abstract base class for architecture artifact exporters."""
    @abstractmethod
    def export(self, data: dict, output_dir: str, project_name: str):
        pass

class JsonExporter(BaseExporter):
    """Exports structured JSON artifacts."""
    def export(self, data: dict, output_dir: str, project_name: str):
        generated = []
        # We can split the massive JSON into separate files for easy downstream use
        artifacts = {
            "architecture_summary.json": data.get("architecture_summary", {}),
            "technology_stack.json": data.get("technology_stack", {}),
            "architecture_decisions.json": data.get("architecture_decisions", []),
            "risks_and_mitigations.json": data.get("technical_risks", []),
            "developer_handoff.json": data.get("handoff_packages", {}).get("developer", {}),
            "qa_handoff.json": data.get("handoff_packages", {}).get("qa", {}),
            "pm_handoff.json": data.get("handoff_packages", {}).get("pm", {})
        }
        for filename, content in artifacts.items():
            path = os.path.join(output_dir, filename)
            with open(path, "w", encoding="utf-8") as f:
                json.dump(content, f, indent=2, ensure_ascii=False)
            generated.append(path)
        print("JSON artifacts exported.")
        return generated

def clean_mermaid(text: str) -> str:
    text = str(text).strip()
    if text.startswith("```mermaid"): text = text[10:]
    elif text.startswith("```"): text = text[3:]
    if text.endswith("```"): text = text[:-3]
    return text.strip()

class MermaidExporter(BaseExporter):
    """Exports raw Mermaid diagrams."""
    def export(self, data: dict, output_dir: str, project_name: str):
        generated = []
        arch_mmd = clean_mermaid(data.get("mermaid_architecture", ""))
        erd_mmd = clean_mermaid(data.get("mermaid_erd", ""))
        
        if arch_mmd:
            path = os.path.join(output_dir, "architecture_diagram.mmd")
            with open(path, "w", encoding="utf-8") as f:
                f.write(arch_mmd)
            generated.append(path)
        if erd_mmd:
            path = os.path.join(output_dir, "database_erd.mmd")
            with open(path, "w", encoding="utf-8") as f:
                f.write(erd_mmd)
            generated.append(path)
        print("Mermaid artifacts exported.")
        return generated

class MarkdownExporter(BaseExporter):
    """Exports a readable Markdown architecture document."""
    def export(self, data: dict, output_dir: str, project_name: str):
        md_content = f"# Architecture Document: {project_name}\n\n"
        md_content += f"**Version:** {data.get('version', '1.0')}\n\n"
        
        summary = data.get("architecture_summary", {})
        md_content += f"## Executive Summary\n{summary.get('overview', '')}\n\n"
        
        tech = data.get("technology_stack", {})
        md_content += "## Technology Stack\n"
        md_content += f"- **Frontend:** {tech.get('frontend', '')}\n"
        md_content += f"- **Backend:** {tech.get('backend', '')}\n"
        md_content += f"- **Database:** {tech.get('database', '')}\n"
        md_content += f"- **Infrastructure:** {tech.get('infrastructure', '')}\n\n"
        
        md_content += "## System Components\n"
        for comp in data.get("system_components", []):
            md_content += f"### {comp.get('name', 'Unknown')}\n"
            md_content += f"{comp.get('description', '')}\n\n"
            
        md_content += "## Architecture Decisions (ADRs)\n"
        for adr in data.get("architecture_decisions", []):
            md_content += f"**{adr.get('title', 'ADR')}**\n"
            md_content += f"> Status: {adr.get('status', 'Proposed')}\n>\n"
            md_content += f"> Context: {adr.get('context', '')}\n>\n"
            md_content += f"> Decision: {adr.get('decision', '')}\n\n"
            
        path = os.path.join(output_dir, "architecture.md")
        with open(path, "w", encoding="utf-8") as f:
            f.write(md_content)
        print("Markdown artifact exported.")
        return [path]

class DocxExporter(BaseExporter):
    """Exports a professional DOCX architecture report."""
    def export(self, data: dict, output_dir: str, project_name: str):
        doc = Document()
        
        # Styles
        style_h1 = doc.styles['Heading 1']
        font_h1 = style_h1.font
        font_h1.name = 'Calibri'
        font_h1.size = Pt(16)
        font_h1.bold = True
        font_h1.color.rgb = RGBColor(0x1F, 0x38, 0x64)
        
        style_h2 = doc.styles['Heading 2']
        font_h2 = style_h2.font
        font_h2.name = 'Calibri'
        font_h2.size = Pt(13)
        font_h2.bold = True
        font_h2.color.rgb = RGBColor(0x2F, 0x2F, 0x2F)
        
        style_normal = doc.styles['Normal']
        font_normal = style_normal.font
        font_normal.name = 'Calibri'
        font_normal.size = Pt(11)
        font_normal.color.rgb = RGBColor(0x00, 0x00, 0x00)
        
        # Title Page
        title = doc.add_paragraph("Technical Solution Architecture")
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.runs[0].font.size = Pt(24)
        title.runs[0].bold = True
        
        doc.add_paragraph(f"Project: {project_name}")
        doc.add_paragraph(f"Date: {datetime.date.today().strftime('%Y-%m-%d')}")
        doc.add_paragraph(f"Version: {data.get('version', '1.0')}")
        
        doc.add_page_break()
        
        # Content
        doc.add_heading("1. Executive Summary", level=1)
        summary = data.get("architecture_summary", {})
        doc.add_paragraph(summary.get("overview", "No summary provided."))
        
        doc.add_heading("2. Technology Stack", level=1)
        tech = data.get("technology_stack", {})
        doc.add_paragraph(f"Frontend: {tech.get('frontend', 'N/A')}", style='List Bullet')
        doc.add_paragraph(f"Backend: {tech.get('backend', 'N/A')}", style='List Bullet')
        doc.add_paragraph(f"Database: {tech.get('database', 'N/A')}", style='List Bullet')
        doc.add_paragraph(f"Infrastructure: {tech.get('infrastructure', 'N/A')}", style='List Bullet')
        
        doc.add_heading("3. System Components", level=1)
        for comp in data.get("system_components", []):
            doc.add_heading(comp.get("name", "Unknown"), level=2)
            doc.add_paragraph(comp.get("description", ""))
            
        doc.add_heading("4. Scalability & Security", level=1)
        doc.add_heading("Scalability Strategy", level=2)
        doc.add_paragraph(str(data.get("scalability_strategy", {}).get("overview", "N/A")))
        doc.add_heading("Security Architecture", level=2)
        doc.add_paragraph(str(data.get("security_architecture", {}).get("authentication", "N/A")))
        
        doc.add_heading("5. Architecture Decisions (ADRs)", level=1)
        for adr in data.get("architecture_decisions", []):
            doc.add_heading(adr.get("title", "ADR"), level=2)
            doc.add_paragraph(f"Context: {adr.get('context', '')}")
            doc.add_paragraph(f"Decision: {adr.get('decision', '')}")
            
        doc_path = os.path.join(output_dir, "Architecture_Report.docx")
        doc.save(doc_path)
        print(f"DOCX artifact exported: {doc_path}")
        return [doc_path]

class ArtifactExporterFramework:
    """Manages generation of exportable artifacts without modifying core agent logic."""
    def __init__(self, output_dir="outputs"):
        self.output_dir = output_dir
        self.exporters = []
        os.makedirs(self.output_dir, exist_ok=True)
        
    def register_exporter(self, exporter: BaseExporter):
        self.exporters.append(exporter)
        
    def run_all(self, data: dict, project_name: str):
        all_generated = []
        for exporter in self.exporters:
            try:
                paths = exporter.export(data, self.output_dir, project_name)
                if paths:
                    all_generated.extend(paths)
            except Exception as e:
                print(f"[Exporter Error] {exporter.__class__.__name__}: {e}")
        return all_generated

# ==========================================
# CORE AGENT LOGIC
# ==========================================

def post_approval(data: dict, context_manager):
    """
    Hook called by the agent runner after human approval.
    Responsible for generating the artifacts using the Exporter Framework.
    """
    context = context_manager.get_context()
    supervisor_output = context.get("Supervisor", {})
    project_name = "Unknown Project"
    if isinstance(supervisor_output, dict):
        project_name = supervisor_output.get("project_name", "Unknown Project")
        
    print("Running Architect post-approval exporters...")
    out_dir = os.path.join("outputs", "Architect")
    framework = ArtifactExporterFramework(output_dir=out_dir)
    
    # Phase 1 Exporters (PDF/PowerPoint can be added to this framework later)
    framework.register_exporter(JsonExporter())
    framework.register_exporter(MermaidExporter())
    framework.register_exporter(MarkdownExporter())
    framework.register_exporter(DocxExporter())
    
    return framework.run_all(data, project_name)

from core.context_compressor import compress_context_for_agent

def run(context_manager, correction: str = None) -> dict:
    """
    Runs the Architect agent.
    Transforms BA outputs into a complete technical solution architecture.
    """
    raw_context = context_manager.get_context()
    context = compress_context_for_agent("Architect", raw_context)
    
    user_brief = context.get("USER_BRIEF", "")
    ba_output = context.get("BA", {})
    architect_output = raw_context.get("Architect", {}) # Fetch previous iteration if editing
    
    prompt_template = """
You are the Expert AI SDLC Solution Architect for ARIA, a multi-agent SDLC platform.
Your responsibility is to transform approved Business Analysis requirements into a production-ready, scalable technical solution architecture.

INPUTS:
=========================================
USER BRIEF:
{user_brief}

APPROVED BUSINESS ANALYSIS:
{ba_output}

PREVIOUS ARCHITECTURE ITERATION (if any):
{previous_arch}

HUMAN FEEDBACK / CORRECTION:
{human_correction}
=========================================

YOUR TASK:
Generate a complete technical solution architecture responding to the business requirements.
You MUST output ONLY valid JSON.

VERSIONING & FEEDBACK:
- If HUMAN FEEDBACK is provided, compare it against the PREVIOUS ARCHITECTURE ITERATION.
- Perform an impact analysis, update ONLY the affected components, preserve unaffected sections, and increment the "version" (e.g., from "1.0" to "1.1").
- If no previous architecture exists, start at "version": "1.0".

DIAGRAM REQUIREMENTS (Mermaid):
- "mermaid_architecture": Must be a valid `graph TD` or `graph LR`. 
  Use standard industry professional shapes and colors via `classDef`. 
  Example:
  classDef frontend fill:#e1f5fe,stroke:#0288d1,stroke-width:2px,color:#000;
  classDef backend fill:#e8f5e9,stroke:#388e3c,stroke-width:2px,color:#000;
  classDef database fill:#fff3e0,stroke:#f57c00,stroke-width:2px,color:#000;
  classDef external fill:#f3e5f5,stroke:#7b1fa2,stroke-width:2px,color:#000;
  Web_App[React Web App]:::frontend --> API_Gateway[FastAPI Gateway]:::backend
- "mermaid_erd": Must be a valid `erDiagram`. Ensure standard entity relationships. Do not mix `erDiagram` syntax with `classDiagram` syntax (e.g., do not use `class` inside an `erDiagram`). Make sure all relationships are complete before closing the string.

OUTPUT SCHEMA (Must match exactly):
{{
  "version": "1.0",
  "architecture_summary": {{
    "overview": "High level description"
  }},
  "technology_stack": {{
    "frontend": "",
    "backend": "",
    "database": "",
    "infrastructure": "",
    "rationale": ""
  }},
  "system_components": [
    {{
      "name": "Component Name",
      "description": "Component Role",
      "technologies": []
    }}
  ],
  "service_boundaries": [
    {{
      "service_name": "",
      "responsibilities": []
    }}
  ],
  "database_design": {{
    "primary_datastore": "",
    "caching_strategy": "",
    "key_entities": []
  }},
  "api_architecture": [
    {{
      "endpoint_group": "",
      "protocol": "REST/GraphQL/gRPC",
      "description": ""
    }}
  ],
  "security_architecture": {{
    "authentication": "",
    "authorization": "",
    "data_protection": ""
  }},
  "scalability_strategy": {{
    "overview": "",
    "horizontal_scaling": "",
    "bottlenecks": []
  }},
  "technical_risks": [
    {{
      "risk": "",
      "impact": "High/Medium/Low",
      "mitigation": ""
    }}
  ],
  "architecture_decisions": [
    {{
      "title": "ADR-001: Selection of Database",
      "status": "Accepted",
      "context": "Why we made this choice...",
      "decision": "What was chosen..."
    }}
  ],
  "mermaid_architecture": "graph TD\\n...",
  "mermaid_erd": "erDiagram\\n...",
  "handoff_packages": {{
    "developer": {{
      "focus": "Key technical patterns to follow...",
      "setup_instructions": ""
    }},
    "qa": {{
      "testable_boundaries": "Where QA should focus...",
      "performance_targets": ""
    }},
    "pm": {{
      "milestone_recommendations": []
    }}
  }},
  "confidence_score": 95,
  "confidence_reasoning": "Explanation of score based on requirement coverage and technical completeness."
}}

CRITICAL RULES:
- Output ONLY raw valid JSON. No markdown backticks outside of the main JSON structure.
- The Mermaid strings must use properly escaped newlines (\\n) and quotes if necessary so the JSON remains valid.
- The architecture must directly address all constraints, assumptions, functional and non-functional requirements from the BA output.
"""

    human_correction_text = f"Correction requested: {correction}" if correction else "None."
    
    prompt = prompt_template.format(
        user_brief=user_brief,
        ba_output=json.dumps(ba_output, indent=2) if ba_output else "No BA output found.",
        previous_arch=json.dumps(architect_output, indent=2) if architect_output else "No previous iteration.",
        human_correction=human_correction_text
    )
    
    print("Calling LLM for Solution Architecture...")
    # Using 8000 max_tokens to accommodate the massive JSON structure and mermaid diagrams
    response_text = call_llm(prompt, agent_name="Architect", max_tokens=8000)
    data = parse_json_from_llm(response_text)
    
    if not data:
        print("Error: LLM returned invalid JSON for Architect.")
        return None
        
    return data
